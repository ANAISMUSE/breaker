from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document

from src.evaluation.index_pipeline import EvaluationResultV2


def export_word_report(
    result: EvaluationResultV2,
    ladder_plan: list[dict],
    df: pd.DataFrame,
    output_dir: str = "outputs",
    strategy_summary: str | None = None,
    platform_placeholder: str | None = None,
) -> str:
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    filename = f"jianping_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = Path(output_dir) / filename

    doc = Document()
    doc.add_heading("茧评 - 信息茧房评估报告（PDF 技术方案口径）", level=1)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"匿名用户：{df['user_id'].iloc[0] if not df.empty and 'user_id' in df.columns else 'unknown'}")

    doc.add_heading("一、四维度多样性得分（1–10，越高越好）", level=2)
    doc.add_paragraph(f"S1 内容类型多样性：{result.s1_content_diversity}")
    doc.add_paragraph(f"S2 跨领域曝光：{result.s2_cross_domain}")
    doc.add_paragraph(f"S3 观点立场多样性：{result.s3_stance_diversity}")
    doc.add_paragraph(f"S4 认知覆盖：{result.s4_cognitive_coverage}")
    doc.add_paragraph(f"综合茧房指数 C（0–10，越高越严重）：{result.cocoon_index}")
    doc.add_paragraph(f"评估模式：{result.mode}")

    doc.add_heading("二、阶梯式破茧计划", level=2)
    for step in ladder_plan:
        doc.add_paragraph(f"{step['level']} | 主题：{step['topic']} | 理由：{step['reason']}")

    if strategy_summary:
        doc.add_heading("三、多策略模拟摘要", level=2)
        doc.add_paragraph(strategy_summary)

    doc.add_heading("四、平台画像与多平台对比（占位）", level=2)
    doc.add_paragraph(
        platform_placeholder
        or "平台级画像与多平台横向对比依赖合规数据源，当前版本未启用自动采集；可在取得合法数据授权后接入统一指标框架。"
    )

    doc.add_heading("五、合规说明", level=2)
    doc.add_paragraph(
        "本报告基于用户主动提供的导出/演示数据，在本地处理；已执行基础脱敏。外接爬虫或自动化采集须自行确保合法授权，本项目默认不启用高风险反爬能力。"
    )
    doc.save(output_path)
    return str(output_path)
